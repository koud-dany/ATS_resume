from pyresparser import ResumeParser
import os 
import nltk
import pandas as pd
import numpy as np
import re 
import string
from nltk.corpus import stopwords
from nltk.stem import LancasterStemmer
from nltk.tokenize import word_tokenize
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from docx import Document

import PyPDF2
import docx
from flask import Flask, request, render_template, url_for, redirect

import spacy
nlp = spacy.load("en_core_web_sm")
nltk.download('punkt')
nltk.download('stopwords')

app = Flask(__name__)

def clean_text(text):
    txt = re.sub("[^a-zA-Z]",  # Search for all non-letters
                          " ",          # Replace all non-letters with spaces
                          str(text.lower()))
    txt_tokens=[word for word in word_tokenize(txt)if word not in string.punctuation]

     #Import the english stop words list from NLTK
    stm= LancasterStemmer()
    txt=' '.join([stm.stem(word) for word in txt_tokens if word not in stopwords.words('english') ])
  
    return txt


def cosim(x, y):

    corpus = [clean_text(x),clean_text(y)]
    count_vect = CountVectorizer()
    X_train_counts = count_vect.fit_transform(corpus)
    cosm= round(cosine_similarity(X_train_counts)[1][0] * 100,2)
    return cosm

def data(path, text):
  file={}
  # Open the PDF file
  pdf_file = open(path, 'rb')
  pdf_reader = PyPDF2.PdfReader(pdf_file)

  # Create a new Word document
  doc = docx.Document()

  # Loop through each page of the PDF file
  for page_num in range(len(pdf_reader.pages)):
      # Get the text from the page
      page = pdf_reader.pages[page_num]
      txt = page.extract_text()
      
      # Add the text to the Word document
      doc.add_paragraph(txt.lower())

  # Save the Word document
  doc.save("docs/Eric_D_fk.docx")
  resume = ResumeParser('docs/Eric_D_fk.docx').get_extracted_data()
  #Job description 
  fl= nlp(text.lower())
  js= fl.text
  # Create a new Word document
  doc_= docx.Document()
  doc_.add_paragraph(js)
  doc_.save("docs/text.docx")
  job = ResumeParser("docs/text.docx").get_extracted_data()
  file['resume']= resume
  file['job']= job

  return file

@app.route('/')
def index():
    return render_template('index.html')


@app.route('/process_text', methods=['POST'])
def process_text():
    if request.method == 'POST':

        f=request.files['file']
        path=f"docs/{f.filename}"  
        f.save(path)
        text1 = request.form['text1']
        text2= request.form['text2']
        ratio=cosim(text1, text2)
        rslt= str(ratio)

        rst= data(path, text2)
        rst['rate']=rslt
        
        resume= rst['resume']
        job=rst['job']
        relev= rst['rate']
        dict_={}
        missing_skill=[]
        for skill in job['skills']:
          if skill not in resume['skills']:
            missing_skill.append(skill)

        count=0
        matching_skill=[]
        for skill in job['skills']:
          if skill in resume['skills']:
            matching_skill.append(skill)
            count += 1

        important_keyword= job['skills']

        job_ratio= round(count / len(important_keyword) * 100,2)
        # resume_ratio= round(count / len(resume['skills']) * 100,2)
        match_df= pd.DataFrame({'matching skills': matching_skill})
        missing_df= pd.DataFrame({'missing skill/keyword':missing_skill})
        keyword= pd.DataFrame({'important keyword':job['skills']})
        txt1=' '.join([word for word in resume['skills']])
        txt2=' '.join([word for word in job['skills']])
        count_vect = CountVectorizer()
        corpus = [txt1,txt2]
        X_counts = count_vect.fit_transform(corpus)
        df= pd.DataFrame(X_counts.toarray(),columns=count_vect.get_feature_names_out(),index=['resume','job description']).T.reset_index().rename(columns={'index': 'skills'})
        skillre=cosim(txt1,txt2)
        dict_['number_match']= str(count)
        dict_['relevance']= relev
        dict_['skill_matching_score']= str(skillre)
        dict_['ratio']= str(job_ratio)
        dict_['total_skill']= str(len(important_keyword) )

        return render_template('result.html', dict_= dict_, table1=[df.to_html()],table2=[match_df.to_html()],table3=[missing_df.to_html()],table4=[keyword.to_html()] )

if __name__ == "__main__":
    app.run(debug= True)

